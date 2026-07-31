"""Format-specific parsers — PDF, DOCX, TXT, MD, HTML, CSV."""

from __future__ import annotations

import csv
import re
from abc import ABC, abstractmethod
from typing import TYPE_CHECKING

from bs4 import BeautifulSoup

if TYPE_CHECKING:
    from pathlib import Path

from rag_pipeline.data.models import ParsedDocument, Section


class BaseParser(ABC):
    """Abstract base for all parsers."""

    @abstractmethod
    def parse(self, file_path: Path) -> ParsedDocument:
        ...


class PDFParser(BaseParser):
    """PDF parser using pdfplumber."""

    def parse(self, file_path: Path) -> ParsedDocument:
        import pdfplumber

        text_parts: list[str] = []
        tables: list[dict] = []
        metadata: dict = {}

        with pdfplumber.open(file_path) as pdf:
            metadata = pdf.metadata or {}
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)

                for table in page.extract_tables():
                    tables.append({
                        "page": i + 1,
                        "data": table,
                        "as_text": self._table_to_text(table),
                    })

        return ParsedDocument(
            content="\n\n".join(text_parts),
            metadata={**metadata, "format": "pdf"},
            tables=tables,
        )

    @staticmethod
    def _table_to_text(table: list[list]) -> str:
        if not table:
            return ""
        headers = [str(c or "") for c in table[0]]
        rows = [[str(c or "") for c in row] for row in table[1:]]
        lines = [" | ".join(headers), "-" * 40]
        lines.extend(" | ".join(row) for row in rows)
        return "\n".join(lines)


class DOCXParser(BaseParser):
    """DOCX parser using python-docx."""

    def parse(self, file_path: Path) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        text_parts: list[str] = []
        tables: list[dict] = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for i, table in enumerate(doc.tables):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append({
                "index": i,
                "data": rows,
                "as_text": self._table_to_text(rows),
            })

        return ParsedDocument(
            content="\n\n".join(text_parts),
            metadata={"title": doc.core_properties.title or "", "format": "docx"},
            tables=tables,
        )

    @staticmethod
    def _table_to_text(table: list[list]) -> str:
        if not table:
            return ""
        lines = [" | ".join(table[0]), "-" * 40]
        lines.extend(" | ".join(row) for row in table[1:])
        return "\n".join(lines)


class TXTParser(BaseParser):
    """Plain text parser with UTF-8 BOM handling."""

    def parse(self, file_path: Path) -> ParsedDocument:
        content = file_path.read_text(encoding="utf-8-sig")  # strips BOM
        return ParsedDocument(content=content, metadata={"format": "txt"})


class MarkdownParser(BaseParser):
    """Markdown parser — preserves heading structure."""

    HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

    def parse(self, file_path: Path) -> ParsedDocument:
        content = file_path.read_text(encoding="utf-8")

        sections = []
        for match in self.HEADING_PATTERN.finditer(content):
            level = len(match.group(1))
            title = match.group(2)
            sections.append(Section(level=level, title=title, offset=match.start()))

        return ParsedDocument(
            content=content,
            metadata={"format": "markdown"},
            sections=sections,
        )


class HTMLParser(BaseParser):
    """HTML parser — strips scripts, styles, nav, footer."""

    def parse(self, file_path: Path) -> ParsedDocument:
        raw = file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(raw, "html.parser")

        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        title_tag = soup.find("title")
        title = title_tag.get_text() if title_tag else ""

        return ParsedDocument(
            content=text,
            metadata={"title": title, "format": "html"},
        )


class CSVParser(BaseParser):
    """CSV parser — each row becomes a text block."""

    def parse(self, file_path: Path) -> ParsedDocument:
        rows: list[str] = []
        with file_path.open(newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                if text.strip():
                    rows.append(text)

        return ParsedDocument(
            content="\n\n".join(rows),
            metadata={"format": "csv", "row_count": len(rows)},
        )


# ---- Registry ---------------------------------------------------------------

PARSERS: dict[str, BaseParser] = {
    ".pdf": PDFParser(),
    ".docx": DOCXParser(),
    ".txt": TXTParser(),
    ".md": MarkdownParser(),
    ".html": HTMLParser(),
    ".htm": HTMLParser(),
    ".csv": CSVParser(),
}


def get_parser(file_path: Path) -> BaseParser:
    """Return the parser for a given file extension."""
    suffix = file_path.suffix.lower()
    if suffix not in PARSERS:
        raise ValueError(f"No parser for extension: {suffix}")
    return PARSERS[suffix]
